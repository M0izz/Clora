"""
Document ingestion interface for PDF, DOCX, and text formats.
Section-aware, table-aware parser.
"""

import os
import re
from typing import List, Optional

from pydantic import BaseModel, Field

EQUIPMENT_ID_REGEX = re.compile(
    r"\b(?:P|K|E|T|TK|V|C|R|HEX|MOV|FCV|PT|TT|LT|FT|D)-\d{2,4}[A-Z]?\b", re.IGNORECASE
)
SECTION_HEADER_REGEX = re.compile(
    r"^(?:(?:\d+\.|\d+\.\d+|\d+\.\d+\.\d+)\s+[A-Z].*|[A-Z][A-Za-z0-9\s\-_/]{2,40}:|#{1,4}\s+.*|[A-Z\s]{4,40})$"
)


class ParsedSection(BaseModel):
    title: str
    page: int
    content: str
    tables: List[str] = Field(default_factory=list)
    equipment_ids: List[str] = Field(default_factory=list)


class IngestedDocument(BaseModel):
    document_id: str
    document_name: str
    document_type: str = "maintenance_report"
    department: str = "maintenance"
    classification: str = "internal"
    allowed_roles: List[str] = Field(
        default_factory=lambda: ["maintenance_engineer", "supervisor", "operator"]
    )
    timestamp: str = "2026-08-31"
    sections: List[ParsedSection] = Field(default_factory=list)


class DocumentParser:
    """Parses PDF, DOCX, TXT, and Markdown files into structured sections."""

    def parse_file(
        self,
        file_path: str,
        document_id: Optional[str] = None,
        document_type: str = "maintenance_report",
        department: str = "maintenance",
        classification: str = "internal",
        allowed_roles: Optional[List[str]] = None,
        timestamp: str = "2026-08-31",
    ) -> IngestedDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        doc_name = os.path.basename(file_path)
        if not document_id:
            document_id = os.path.splitext(doc_name)[0].lower().replace(" ", "_").replace("-", "_")

        if allowed_roles is None:
            allowed_roles = ["maintenance_engineer", "supervisor", "operator"]

        sections: List[ParsedSection] = []

        if ext == ".pdf":
            sections = self._parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            sections = self._parse_docx(file_path)
        else:
            sections = self._parse_text_file(file_path)

        return IngestedDocument(
            document_id=document_id,
            document_name=doc_name,
            document_type=document_type,
            department=department,
            classification=classification,
            allowed_roles=allowed_roles,
            timestamp=timestamp,
            sections=sections,
        )

    def _parse_pdf(self, file_path: str) -> List[ParsedSection]:
        sections: List[ParsedSection] = []
        try:
            import fitz

            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                sections.extend(self._extract_sections_from_page(text, page_num + 1))
            doc.close()
        except Exception:
            pass
        return sections

    def _parse_docx(self, file_path: str) -> List[ParsedSection]:
        sections: List[ParsedSection] = []
        try:
            import docx

            doc = docx.Document(file_path)
            current_title = "General"
            current_lines = []
            tables_text = []

            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                if p.style.name.startswith("Heading") or SECTION_HEADER_REGEX.match(text):
                    if current_lines:
                        c_str = "\n".join(current_lines)
                        eqs = list(set(EQUIPMENT_ID_REGEX.findall(c_str)))
                        sections.append(
                            ParsedSection(
                                title=current_title,
                                page=1,
                                content=c_str,
                                tables=tables_text,
                                equipment_ids=eqs,
                            )
                        )
                        current_lines = []
                        tables_text = []
                    current_title = text
                else:
                    current_lines.append(text)

            for table in doc.tables:
                rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                tables_text.append("\n".join(rows))

            if current_lines or tables_text:
                c_str = "\n".join(current_lines)
                eqs = list(set(EQUIPMENT_ID_REGEX.findall(c_str + " " + " ".join(tables_text))))
                sections.append(
                    ParsedSection(
                        title=current_title,
                        page=1,
                        content=c_str,
                        tables=tables_text,
                        equipment_ids=eqs,
                    )
                )
        except Exception:
            pass
        return sections

    def _parse_text_file(self, file_path: str) -> List[ParsedSection]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return self._extract_sections_from_page(text, page=1)

    def _extract_sections_from_page(self, page_text: str, page: int) -> List[ParsedSection]:
        sections: List[ParsedSection] = []
        lines = page_text.splitlines()
        current_title = "General"
        current_lines = []
        tables_text = []

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            if "|" in line_str and line_str.startswith("|") and line_str.endswith("|"):
                tables_text.append(line_str)
                continue
            if SECTION_HEADER_REGEX.match(line_str) and len(line_str) < 60:
                if current_lines:
                    c_str = "\n".join(current_lines)
                    eqs = list(set(EQUIPMENT_ID_REGEX.findall(c_str)))
                    sections.append(
                        ParsedSection(
                            title=current_title,
                            page=page,
                            content=c_str,
                            tables=tables_text,
                            equipment_ids=eqs,
                        )
                    )
                    current_lines = []
                    tables_text = []
                current_title = line_str.lstrip("#").strip().rstrip(":")
            else:
                current_lines.append(line_str)

        if current_lines or tables_text:
            c_str = "\n".join(current_lines)
            eqs = list(set(EQUIPMENT_ID_REGEX.findall(c_str + " " + " ".join(tables_text))))
            sections.append(
                ParsedSection(
                    title=current_title,
                    page=page,
                    content=c_str,
                    tables=tables_text,
                    equipment_ids=eqs,
                )
            )

        return sections
