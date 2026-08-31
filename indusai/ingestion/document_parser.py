"""
INDUSAI-X Intelligent Document Parser
Extracts section hierarchies, tables, page numbers, and equipment identifiers from industrial documents.
"""

import os
import re
from typing import List, Dict, Any, Optional
from indusai.ingestion.schema import ParsedSection, IngestedDocument

EQUIPMENT_ID_REGEX = re.compile(r'\b(?:P|K|E|T|TK|V|C|R|HEX|MOV|FCV|PT|TT|LT|FT|D)-\d{2,4}[A-Z]?\b', re.IGNORECASE)
SECTION_HEADER_REGEX = re.compile(r'^(?:(?:\d+\.|\d+\.\d+|\d+\.\d+\.\d+)\s+[A-Z].*|[A-Z][A-Za-z0-9\s\-_/]{2,40}:|#{1,4}\s+.*|[A-Z\s]{4,40})$')

class DocumentParser:
    """Parses PDF, DOCX, TXT, and Markdown files into structured sections with page and table awareness."""

    def __init__(self):
        pass

    def parse_file(
        self,
        file_path: str,
        document_id: Optional[str] = None,
        document_type: str = "maintenance_report",
        department: str = "maintenance",
        classification: str = "internal",
        allowed_roles: Optional[List[str]] = None,
        timestamp: str = "2026-08-31"
    ) -> IngestedDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        doc_name = os.path.basename(file_path)
        if not document_id:
            document_id = os.path.splitext(doc_name)[0].lower().replace(" ", "_").replace("-", "_")

        if allowed_roles is None:
            allowed_roles = ["maintenance_engineer", "supervisor", "operator"]

        sections: List[ParsedSection] = []

        if ext == ".pdf":
            sections = self._parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            sections = self._parse_docx(file_path)
        else:
            sections = self._parse_text_file(file_path)

        return IngestedDocument(
            document_id=document_id,
            document_name=doc_name,
            document_type=document_type,
            department=department,
            classification=classification,
            allowed_roles=allowed_roles,
            timestamp=timestamp,
            sections=sections
        )

    def _parse_pdf(self, file_path: str) -> List[ParsedSection]:
        sections: List[ParsedSection] = []
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                page_sections = self._extract_sections_from_page(text, page_num + 1)
                sections.extend(page_sections)
            doc.close()
        except ImportError:
            # Fallback if PyMuPDF fails
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    page_sections = self._extract_sections_from_page(text, page_num + 1)
                    sections.extend(page_sections)
            except Exception as e:
                sections.append(ParsedSection(
                    title="Document Content",
                    page=1,
                    content=f"Error reading PDF: {e}",
                    tables=[],
                    equipment_ids=[]
                ))
        return sections

    def _parse_docx(self, file_path: str) -> List[ParsedSection]:
        sections: List[ParsedSection] = []
        try:
            import docx
            doc = docx.Document(file_path)
            current_title = "Document Content"
            current_lines: List[str] = []
            tables_text: List[str] = []

            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                if p.style.name.startswith("Heading") or SECTION_HEADER_REGEX.match(text):
                    if current_lines:
                        content_str = "\n".join(current_lines)
                        eq_ids = list(set(EQUIPMENT_ID_REGEX.findall(content_str)))
                        sections.append(ParsedSection(
                            title=current_title,
                            page=1,
                            content=content_str,
                            tables=tables_text,
                            equipment_ids=eq_ids
                        ))
                        current_lines = []
                        tables_text = []
                    current_title = text
                else:
                    current_lines.append(text)

            # Process tables in DOCX
            for table in doc.tables:
                rows_data = []
                for row in table.rows:
                    rows_data.append(" | ".join(cell.text.strip() for cell in row.cells))
                table_str = "\n".join(rows_data)
                tables_text.append(table_str)

            if current_lines or tables_text:
                content_str = "\n".join(current_lines)
                eq_ids = list(set(EQUIPMENT_ID_REGEX.findall(content_str + " " + " ".join(tables_text))))
                sections.append(ParsedSection(
                    title=current_title,
                    page=1,
                    content=content_str,
                    tables=tables_text,
                    equipment_ids=eq_ids
                ))
        except Exception as e:
            sections.append(ParsedSection(
                title="Document Content",
                page=1,
                content=f"Error reading DOCX: {e}",
                tables=[],
                equipment_ids=[]
            ))
        return sections

    def _parse_text_file(self, file_path: str) -> List[ParsedSection]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return self._extract_sections_from_page(text, page=1)

    def _extract_sections_from_page(self, page_text: str, page: int) -> List[ParsedSection]:
        sections: List[ParsedSection] = []
        lines = page_text.splitlines()
        current_title = "General"
        current_lines: List[str] = []
        tables_text: List[str] = []

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            # Detect markdown table line
            if "|" in line_str and line_str.startswith("|") and line_str.endswith("|"):
                tables_text.append(line_str)
                continue

            # Detect section title
            if SECTION_HEADER_REGEX.match(line_str) and len(line_str) < 60:
                if current_lines:
                    content_str = "\n".join(current_lines)
                    eq_ids = list(set(EQUIPMENT_ID_REGEX.findall(content_str)))
                    sections.append(ParsedSection(
                        title=current_title,
                        page=page,
                        content=content_str,
                        tables=tables_text,
                        equipment_ids=eq_ids
                    ))
                    current_lines = []
                    tables_text = []
                current_title = line_str.lstrip("#").strip().rstrip(":")
            else:
                current_lines.append(line_str)

        if current_lines or tables_text:
            content_str = "\n".join(current_lines)
            eq_ids = list(set(EQUIPMENT_ID_REGEX.findall(content_str + " " + " ".join(tables_text))))
            sections.append(ParsedSection(
                title=current_title,
                page=page,
                content=content_str,
                tables=tables_text,
                equipment_ids=eq_ids
            ))

        return sections
