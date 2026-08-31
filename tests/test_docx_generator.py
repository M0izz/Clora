"""
Tests for Word Deliverable Generator (.docx).
"""

import os
import unittest
from docx import Document
from data_intelligence.docx_generator import ApprovalNoteGenerator
from data_intelligence.models import ApprovalNoteInput, FindingItem


class TestDocxGenerator(unittest.TestCase):

    def setUp(self):
        self.generator = ApprovalNoteGenerator()
        self.output_docx = "tests/test_generated_note.docx"
        if os.path.exists(self.output_docx):
            os.remove(self.output_docx)

    def tearDown(self):
        if os.path.exists(self.output_docx):
            os.remove(self.output_docx)

    def test_generate_approval_note_typed(self):
        note = ApprovalNoteInput(
            note_number="MRPL/TEST/2026/001",
            department="Inspection & Maintenance Dept.",
            date_str="31-Aug-2026",
            subject="Test Emergency Overhaul",
            priority="URGENT",
            author_name="R. Kumar (Sr. Eng.)",
            approver_name="CGM (Technical Services)",
            executive_summary="Vibration alert on Pump P-102A.",
            findings=[
                FindingItem(
                    equipment_tag="P-102A",
                    parameter="DE Vibration",
                    observed_value="7.8 mm/s",
                    threshold_limit="4.5 mm/s",
                    severity="CRITICAL",
                    action_required="Overhaul"
                )
            ],
            risk_assessment="High risk of seal breach.",
            financial_estimate_inr=150000.0,
            recommendation="Approve replacement immediately.",
            output_docx_path=self.output_docx
        )

        res_path = self.generator.generate(note)
        self.assertTrue(os.path.exists(res_path))

        # Inspect generated docx
        doc = Document(res_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("MANGALORE REFINERY", full_text)
        self.assertIn("SUBJECT: Test Emergency Overhaul", full_text)
        self.assertIn("INR 150,000.00", full_text)

        # Check table count (Metadata, Summary, Findings, Sign-off)
        self.assertEqual(len(doc.tables), 4)

    def test_generate_approval_note_from_flat_dict(self):
        flat_payload = {
            "note_number": "MRPL/FLAT/2026/002",
            "department": "Safety Dept.",
            "date_str": "31-Aug-2026",
            "subject": "Flat Payload Test",
            "priority": "HIGH",
            "author_name": "Safety Officer",
            "approver_name": "GM Safety",
            "executive_summary": "Gas detector calibration check.",
            "findings_summary": "Detector GD-101 sensor drifted; recalibration needed.",
            "risk_assessment": "Low risk under ventilation.",
            "financial_estimate_inr": 25000.0,
            "recommendation": "Calibrate sensor.",
            "output_docx_path": self.output_docx
        }
        res_path = self.generator.generate(flat_payload)
        self.assertTrue(os.path.exists(res_path))
        doc = Document(res_path)
        self.assertTrue(len(doc.paragraphs) > 0)


if __name__ == "__main__":
    unittest.main()
