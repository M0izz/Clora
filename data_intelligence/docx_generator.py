"""
MRPL Boardroom Deliverable & Word Approval Note Generator.
INDUSAI-X / SIH Problem Statement 26117 (MRPL)
Member 6: Data Intelligence + Knowledge Graph + Security Engineer

Builds formatted, executive-ready Word (.docx) approval notes from structured
agent findings (ApprovalNoteInput or flat LLM tool payloads).
"""

import os
from typing import Dict, Any, Union
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from .models import ApprovalNoteInput, FindingItem


def set_cell_background(cell, hex_color: str):
    """Sets background shading for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


class ApprovalNoteGenerator:
    """Generates official MRPL executive approval notes in Word (.docx) format."""

    @staticmethod
    def _parse_payload(payload: Union[ApprovalNoteInput, Dict[str, Any]]) -> ApprovalNoteInput:
        """Adapts either typed ApprovalNoteInput or flat dictionary from Member 3 tool call."""
        if isinstance(payload, ApprovalNoteInput):
            return payload

        # Flat dictionary adapter
        findings_raw = payload.get("findings", [])
        if not findings_raw and "findings_summary" in payload:
            # Construct single summary finding item if plain text passed
            findings_raw = [
                FindingItem(
                    equipment_tag="Refinery Unit",
                    parameter="General Health Inspection",
                    observed_value="Abnormal",
                    threshold_limit="Normal",
                    severity=payload.get("priority", "HIGH"),
                    action_required=payload.get("findings_summary", "Review required")
                )
            ]

        findings = []
        for f in findings_raw:
            if isinstance(f, FindingItem):
                findings.append(f)
            elif isinstance(f, dict):
                findings.append(FindingItem(**f))

        return ApprovalNoteInput(
            note_number=payload.get("note_number", "MRPL/MAINT/2026/001"),
            department=payload.get("department", "Inspection & Maintenance Dept."),
            date_str=payload.get("date_str", "31-Aug-2026"),
            subject=payload.get("subject", "Technical Approval Note"),
            priority=payload.get("priority", "HIGH"),
            author_name=payload.get("author_name", "Maintenance Engineer"),
            approver_name=payload.get("approver_name", "Chief General Manager (TS)"),
            executive_summary=payload.get("executive_summary", ""),
            findings=findings,
            risk_assessment=payload.get("risk_assessment", ""),
            financial_estimate_inr=float(payload.get("financial_estimate_inr", 0.0)),
            recommendation=payload.get("recommendation", ""),
            output_docx_path=payload.get("output_docx_path", "approval_note.docx")
        )

    def generate(self, payload: Union[ApprovalNoteInput, Dict[str, Any]], output_path: str = None) -> str:
        """Builds and saves the formatted .docx approval note."""
        note = self._parse_payload(payload)
        target_path = output_path or note.output_docx_path
        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)

        doc = Document()

        # Page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # 1. Official Header
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_run1 = header_p.add_run("MANGALORE REFINERY AND PETROCHEMICALS LIMITED\n")
        h_run1.bold = True
        h_run1.font.size = Pt(14)
        h_run1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)  # MRPL Navy

        h_run2 = header_p.add_run("CONFIDENTIAL INTERNAL APPROVAL NOTE & TECHNICAL SANCTION")
        h_run2.bold = True
        h_run2.font.size = Pt(10)
        h_run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Horizontal separator rule
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # 2. Metadata Grid Table
        meta_table = doc.add_table(rows=3, cols=4)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_data = [
            ("Note Ref No:", note.note_number, "Date:", note.date_str),
            ("Department:", note.department, "Priority Level:", note.priority),
            ("Originator:", note.author_name, "Approver:", note.approver_name)
        ]

        for r_idx, row in enumerate(meta_data):
            for c_idx in range(4):
                cell = meta_table.cell(r_idx, c_idx)
                cell.text = row[c_idx]
                set_cell_background(cell, "F1F5F9" if c_idx % 2 == 0 else "FFFFFF")
                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                if c_idx % 2 == 0:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # 3. Subject Line
        subj_p = doc.add_paragraph()
        subj_p.paragraph_format.space_before = Pt(6)
        subj_p.paragraph_format.space_after = Pt(8)
        s_lbl = subj_p.add_run("SUBJECT: ")
        s_lbl.bold = True
        s_lbl.font.size = Pt(11)
        s_lbl.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        s_val = subj_p.add_run(note.subject)
        s_val.bold = True
        s_val.font.size = Pt(11)

        # 4. Section 1: Executive Summary Callout Box
        h1 = doc.add_heading("1. Executive Summary", level=2)
        h1.style.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        
        summary_table = doc.add_table(rows=1, cols=1)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        s_cell = summary_table.cell(0, 0)
        set_cell_background(s_cell, "F8FAFC")
        set_cell_margins(s_cell, top=140, bottom=140, left=160, right=160)
        sp = s_cell.paragraphs[0]
        s_run = sp.add_run(note.executive_summary)
        s_run.font.size = Pt(10)
        s_run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # 5. Section 2: Equipment Findings Table
        if note.findings:
            h2 = doc.add_heading("2. Detailed Equipment Findings & Telemetry", level=2)
            h2.style.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

            findings_table = doc.add_table(rows=len(note.findings) + 1, cols=6)
            findings_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = ["Equipment Tag", "Parameter", "Observed", "Threshold", "Severity", "Action Required"]
            for c_idx, h_text in enumerate(headers):
                cell = findings_table.cell(0, c_idx)
                cell.text = h_text
                set_cell_background(cell, "003366")
                set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(8.5)
                    p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for r_idx, f in enumerate(note.findings):
                row_cells = [
                    f.equipment_tag,
                    f.parameter,
                    f.observed_value,
                    f.threshold_limit,
                    f.severity,
                    f.action_required
                ]
                bg_color = "FFF1F2" if f.severity.upper() == "CRITICAL" else ("FEFCE8" if f.severity.upper() == "WARNING" else "FFFFFF")
                for c_idx, val in enumerate(row_cells):
                    cell = findings_table.cell(r_idx + 1, c_idx)
                    cell.text = val
                    set_cell_background(cell, bg_color)
                    set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
                    p = cell.paragraphs[0]
                    if p.runs:
                        p.runs[0].font.size = Pt(8.5)
                        if c_idx == 4:  # Severity column
                            p.runs[0].bold = True
                            if f.severity.upper() == "CRITICAL":
                                p.runs[0].font.color.rgb = RGBColor(0xBE, 0x12, 0x3C)
                            elif f.severity.upper() == "WARNING":
                                p.runs[0].font.color.rgb = RGBColor(0xA1, 0x62, 0x07)

            doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # 6. Section 3: Risk Assessment
        if note.risk_assessment:
            h3 = doc.add_heading("3. Operational Risk & Safety Assessment", level=2)
            h3.style.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            rp = doc.add_paragraph(note.risk_assessment)
            rp.paragraph_format.space_after = Pt(6)

        # 7. Section 4: Budget & Financial Sanction
        if note.financial_estimate_inr > 0:
            h4 = doc.add_heading("4. Financial Budget Estimate", level=2)
            h4.style.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            fp = doc.add_paragraph()
            fp.add_run("Estimated Financial Sanction Required: ").bold = True
            fp.add_run(f"INR {note.financial_estimate_inr:,.2f} ").bold = True
            fp.add_run("(Inclusive of required OEM spare parts, precision balancing, and overhaul labor).")
            fp.paragraph_format.space_after = Pt(6)

        # 8. Section 5: Recommendations
        h5 = doc.add_heading("5. Recommendation & Sanction Request", level=2)
        h5.style.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        recomp = doc.add_paragraph(note.recommendation)
        recomp.paragraph_format.space_after = Pt(16)

        # 9. Official Sign-Off Block
        sign_table = doc.add_table(rows=2, cols=3)
        sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sign_data = [
            ("Initiated By:", "Reviewed By:", "Approved By:"),
            (
                f"{note.author_name}\nSr. Maintenance Engineer",
                "Head of Maintenance (CDU-1)\nMRPL Technical Services",
                f"{note.approver_name}\nChief General Manager"
            )
        ]
        for r_idx, s_row in enumerate(sign_data):
            for c_idx, text in enumerate(s_row):
                cell = sign_table.cell(r_idx, c_idx)
                cell.text = text
                set_cell_background(cell, "F8FAFC" if r_idx == 0 else "FFFFFF")
                set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].font.size = Pt(8.5)
                    if r_idx == 0:
                        p.runs[0].bold = True

        doc.save(target_path)
        return target_path
